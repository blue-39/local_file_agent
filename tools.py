import ast
import difflib
import json
import operator
import re
from collections import Counter
from pathlib import Path
from typing import Iterable
from datetime import datetime
import shutil
import tomllib


TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".toml",
    ".csv", ".log", ".ini", ".cfg", ".html", ".css", ".js",
    ".ts", ".tsx", ".jsx", ".java", ".cpp", ".c", ".h", ".hpp",
    ".go", ".rs", ".sh", ".bat", ".sql", ".xml"
}

DOC_EXTENSIONS = {
    ".pdf", ".docx"
}

SUPPORTED_READ_EXTENSIONS = TEXT_EXTENSIONS | DOC_EXTENSIONS

MAX_FILE_CHARS = 10000
MAX_TOTAL_CHARS = 30000
MAX_SEARCH_RESULTS = 80
MAX_LIST_ITEMS = 300
MAX_TREE_ITEMS = 500

SKIP_DIRS = {
    ".git", ".hg", ".svn", "__pycache__", ".pytest_cache", ".mypy_cache",
    ".ruff_cache", ".idea", ".vscode", "node_modules", ".venv", "venv",
    "dist", "build", ".next", ".turbo",
}

CODE_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".cpp", ".c", ".h",
    ".hpp", ".go", ".rs", ".sh", ".bat", ".sql", ".html", ".css",
}

_CALC_OPERATORS = {
    ast.Add: operator.add,
    ast.Sub: operator.sub,
    ast.Mult: operator.mul,
    ast.Div: operator.truediv,
    ast.FloorDiv: operator.floordiv,
    ast.Mod: operator.mod,
    ast.Pow: operator.pow,
    ast.USub: operator.neg,
    ast.UAdd: operator.pos,
}


def _resolve_root(root_dir: str) -> Path:
    root = Path(root_dir).expanduser().resolve()

    if not root.exists():
        raise ValueError(f"根目录不存在：{root_dir}")

    if not root.is_dir():
        raise ValueError(f"不是目录：{root_dir}")

    return root


def _safe_path(root_dir: str, path: str = ".") -> Path:
    root = _resolve_root(root_dir)
    target = (root / path).expanduser().resolve()

    try:
        target.relative_to(root)
    except ValueError:
        raise ValueError("非法路径：只能访问用户指定的根目录内部。")

    return target


def _is_supported_read_file(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_READ_EXTENSIONS or path.suffix == ""


def _is_text_file(path: Path) -> bool:
    return path.suffix.lower() in TEXT_EXTENSIONS or path.suffix == ""


def _is_docx_file(path: Path) -> bool:
    return path.suffix.lower() == ".docx"


def _is_supported_write_file(path: Path) -> bool:
    return _is_text_file(path) or _is_docx_file(path)


def _iter_files(root: Path) -> Iterable[Path]:
    if root.is_file():
        yield root
        return

    for path in root.rglob("*"):
        if any(part in SKIP_DIRS for part in path.parts):
            continue

        if path.is_file():
            yield path


def _iter_supported_files(root: Path) -> Iterable[Path]:
    for path in _iter_files(root):
        if _is_supported_read_file(path):
            yield path


def _backup_file(path: Path) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    backup_path = path.with_suffix(path.suffix + f".bak_{timestamp}")
    shutil.copy2(path, backup_path)
    return backup_path


def _add_docx_content(doc, content: str) -> None:
    lines = content.splitlines()

    if not lines:
        doc.add_paragraph("")
        return

    for line in lines:
        doc.add_paragraph(line)


def _write_docx(path: Path, content: str) -> None:
    from docx import Document

    doc = Document()
    _add_docx_content(doc, content)
    doc.save(str(path))


def _append_docx(path: Path, content: str) -> None:
    from docx import Document

    doc = Document(str(path)) if path.exists() else Document()
    _add_docx_content(doc, content)
    doc.save(str(path))


def _replace_docx_text(path: Path, old_text: str, new_text: str) -> bool:
    from docx import Document

    doc = Document(str(path))
    changed = False

    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
            changed = True

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_text, new_text)
                        changed = True

    if changed:
        doc.save(str(path))

    return changed


def _safe_eval_math(node):
    if isinstance(node, ast.Expression):
        return _safe_eval_math(node.body)

    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value

    if isinstance(node, ast.BinOp) and type(node.op) in _CALC_OPERATORS:
        left = _safe_eval_math(node.left)
        right = _safe_eval_math(node.right)

        if isinstance(node.op, ast.Pow) and abs(right) > 10:
            raise ValueError("指数过大，已拒绝计算。")

        return _CALC_OPERATORS[type(node.op)](left, right)

    if isinstance(node, ast.UnaryOp) and type(node.op) in _CALC_OPERATORS:
        value = _safe_eval_math(node.operand)
        return _CALC_OPERATORS[type(node.op)](value)

    raise ValueError("表达式包含不支持的语法。")


def calculator(expression: str) -> str:
    try:
        if len(expression) > 120:
            return "表达式过长，已拒绝计算。"

        allowed_chars = set("0123456789+-*/%(). ")
        if not set(expression) <= allowed_chars:
            return "表达式包含非法字符。"

        tree = ast.parse(expression, mode="eval")
        return str(_safe_eval_math(tree))

    except Exception as e:
        return f"计算失败：{e}"


def list_dir(root_dir: str, path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)

        if not target.exists():
            return f"路径不存在：{path}"

        if not target.is_dir():
            return f"不是目录：{path}"

        root = _resolve_root(root_dir)
        items = []

        for child in sorted(target.iterdir()):
            relative = child.relative_to(root)
            kind = "DIR " if child.is_dir() else "FILE"
            items.append(f"{kind} {relative}")

            if len(items) >= MAX_LIST_ITEMS:
                items.append("[结果过多，已截断]")
                break

        return "\n".join(items) if items else "目录为空。"

    except Exception as e:
        return f"列目录失败：{e}"


def find_files(root_dir: str, pattern: str = "*", path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        candidates = [target] if target.is_file() else target.rglob(pattern)
        results = []

        for file_path in candidates:
            if file_path.is_file():
                results.append(str(file_path.relative_to(root)))

                if len(results) >= MAX_SEARCH_RESULTS:
                    return "\n".join(results) + "\n\n[结果过多，已截断]"

        return "\n".join(results) if results else "没有找到匹配文件。"

    except Exception as e:
        return f"查找文件失败：{e}"


def get_file_info(root_dir: str, path: str) -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        relative = target.relative_to(root)

        if target.is_dir():
            file_count = sum(1 for p in target.rglob("*") if p.is_file())
            dir_count = sum(1 for p in target.rglob("*") if p.is_dir())

            return (
                f"路径：{relative}\n"
                f"类型：目录\n"
                f"子目录数量：{dir_count}\n"
                f"文件数量：{file_count}"
            )

        size = target.stat().st_size

        return (
            f"路径：{relative}\n"
            f"类型：文件\n"
            f"扩展名：{target.suffix or '[无扩展名]'}\n"
            f"大小：{size} bytes"
        )

    except Exception as e:
        return f"获取信息失败：{e}"


def _read_pdf(path: Path, max_chars: int) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks = []

        for i, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(f"[Page {i}]\n{text.strip()}")

            current_len = sum(len(x) for x in chunks)
            if current_len >= max_chars:
                break

        text = "\n\n".join(chunks)

        if not text.strip():
            return "PDF 未提取到文本。可能是扫描版 PDF，需要 OCR。"

        return text[:max_chars] + ("\n\n[PDF 内容过长，已截断]" if len(text) > max_chars else "")

    except Exception as e:
        return f"读取 PDF 失败：{e}"


def _read_docx(path: Path, max_chars: int) -> str:
    try:
        from docx import Document

        doc = Document(str(path))
        chunks = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                chunks.append(text)

        for table_idx, table in enumerate(doc.tables, start=1):
            chunks.append(f"\n[Table {table_idx}]")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    chunks.append(row_text)

        text = "\n".join(chunks)

        if not text.strip():
            return "DOCX 未提取到文本。"

        return text[:max_chars] + ("\n\n[DOCX 内容过长，已截断]" if len(text) > max_chars else "")

    except Exception as e:
        return f"读取 DOCX 失败：{e}"


def _read_text(path: Path, max_chars: int) -> str:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return text[:max_chars] + ("\n\n[文件内容过长，已截断]" if len(text) > max_chars else "")


def read_file(root_dir: str, path: str, max_chars: int = MAX_FILE_CHARS) -> str:
    try:
        target = _safe_path(root_dir, path)

        if not target.exists():
            return f"文件不存在：{path}"

        if not target.is_file():
            return f"不是文件：{path}"

        if not _is_supported_read_file(target):
            return f"暂不支持读取该类型文件：{target.suffix}"

        suffix = target.suffix.lower()

        if suffix == ".pdf":
            content = _read_pdf(target, max_chars)
        elif suffix == ".docx":
            content = _read_docx(target, max_chars)
        else:
            content = _read_text(target, max_chars)

        return f"[FILE: {path}]\n{content}"

    except Exception as e:
        return f"读取文件失败：{e}"


def read_file_chunk(
    root_dir: str,
    path: str,
    start_line: int = 1,
    line_count: int = 120,
) -> str:
    try:
        target = _safe_path(root_dir, path)

        if not target.exists():
            return f"文件不存在：{path}"

        if not target.is_file():
            return f"不是文件：{path}"

        if not _is_text_file(target):
            return f"当前只支持按行读取纯文本文件，不支持：{target.suffix}"

        start_line = max(1, start_line)
        line_count = min(max(1, line_count), 300)
        end_line = start_line + line_count - 1

        selected = []
        with target.open("r", encoding="utf-8", errors="ignore") as f:
            for line_no, line in enumerate(f, start=1):
                if line_no < start_line:
                    continue

                if line_no > end_line:
                    break

                selected.append(f"{line_no}: {line.rstrip()}")

        if not selected:
            return f"文件 {path} 在第 {start_line} 行之后没有内容。"

        return f"[FILE CHUNK: {path}:{start_line}-{end_line}]\n" + "\n".join(selected)

    except Exception as e:
        return f"按行读取文件失败：{e}"


def read_multiple_files(
    root_dir: str,
    paths: list[str],
    max_chars_each: int = 8000,
    max_total_chars: int = MAX_TOTAL_CHARS,
) -> str:
    """
    一次读取多个文件，用于综合分析。
    """
    try:
        outputs = []
        total = 0

        for path in paths:
            content = read_file(root_dir=root_dir, path=path, max_chars=max_chars_each)

            if total + len(content) > max_total_chars:
                remaining = max_total_chars - total
                if remaining <= 0:
                    outputs.append("\n[总内容过长，后续文件未读取]")
                    break

                content = content[:remaining] + "\n\n[总内容过长，当前文件已截断]"

            outputs.append(content)
            total += len(content)

        return "\n\n" + ("\n\n" + "=" * 60 + "\n\n").join(outputs)

    except Exception as e:
        return f"读取多个文件失败：{e}"


def project_overview(root_dir: str, path: str = ".", max_items: int = MAX_TREE_ITEMS) -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        files = []
        dirs = set()
        ext_counter = Counter()
        marker_files = []

        candidates = [target] if target.is_file() else _iter_files(target)

        for file_path in candidates:
            relative = file_path.relative_to(root)
            files.append(relative)
            ext_counter[file_path.suffix.lower() or "[无扩展名]"] += 1

            for parent in relative.parents:
                if str(parent) != ".":
                    dirs.add(parent)

            if file_path.name.lower() in {
                "readme.md", "pyproject.toml", "requirements.txt",
                "package.json", "tsconfig.json", "vite.config.ts",
                "next.config.js", "dockerfile", "compose.yaml",
            }:
                marker_files.append(str(relative))

            if len(files) >= max_items:
                break

        tree_lines = []
        for relative in sorted(files)[:max_items]:
            depth = len(relative.parts) - 1
            tree_lines.append(f"{'  ' * depth}- {relative.name}")

        ext_text = ", ".join(f"{ext}: {count}" for ext, count in ext_counter.most_common(12))
        markers_text = "\n".join(f"- {item}" for item in sorted(marker_files)) or "未发现常见项目标识文件。"
        truncated = "\n[结果过多，已截断]" if len(files) >= max_items else ""

        return (
            f"项目概览：{path}\n"
            f"文件数：{len(files)}\n"
            f"目录数：{len(dirs)}\n"
            f"主要扩展名：{ext_text or '无'}\n\n"
            f"关键文件：\n{markers_text}\n\n"
            f"文件树节选：\n" + ("\n".join(tree_lines) or "无文件") + truncated
        )

    except Exception as e:
        return f"生成项目概览失败：{e}"


def code_stats(root_dir: str, path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        total_files = 0
        total_lines = 0
        ext_stats = Counter()
        largest_files = []

        for file_path in _iter_files(target):
            if file_path.suffix.lower() not in CODE_EXTENSIONS:
                continue

            try:
                line_count = sum(1 for _ in file_path.open("r", encoding="utf-8", errors="ignore"))
            except Exception:
                continue

            total_files += 1
            total_lines += line_count
            ext_stats[file_path.suffix.lower()] += line_count
            largest_files.append((line_count, str(file_path.relative_to(root))))

        largest_files.sort(reverse=True)
        ext_text = "\n".join(f"- {ext}: {lines} 行" for ext, lines in ext_stats.most_common())
        largest_text = "\n".join(f"- {name}: {lines} 行" for lines, name in largest_files[:10])

        return (
            f"代码统计：{path}\n"
            f"代码文件数：{total_files}\n"
            f"总行数：{total_lines}\n\n"
            f"按语言/扩展名统计：\n{ext_text or '无代码文件'}\n\n"
            f"最大文件 Top 10：\n{largest_text or '无代码文件'}"
        )

    except Exception as e:
        return f"统计代码失败：{e}"


def regex_search(
    root_dir: str,
    pattern: str,
    path: str = ".",
    ignore_case: bool = True,
    max_results: int = MAX_SEARCH_RESULTS,
) -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        flags = re.IGNORECASE if ignore_case else 0
        regex = re.compile(pattern, flags)
        results = []
        files = [target] if target.is_file() else _iter_supported_files(target)

        for file_path in files:
            if not _is_text_file(file_path):
                continue

            try:
                lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
            except Exception:
                continue

            for line_no, line in enumerate(lines, start=1):
                if regex.search(line):
                    relative = file_path.relative_to(root)
                    results.append(f"{relative}:{line_no}: {line.strip()}")

                    if len(results) >= max_results:
                        return "\n".join(results) + "\n\n[正则搜索结果过多，已截断]"

        return "\n".join(results) if results else "没有找到匹配内容。"

    except Exception as e:
        return f"正则搜索失败：{e}"


def todo_report(root_dir: str, path: str = ".") -> str:
    try:
        pattern = r"\b(TODO|FIXME|BUG|HACK|XXX)\b[:：]?\s*(.*)"
        result = regex_search(
            root_dir=root_dir,
            pattern=pattern,
            path=path,
            ignore_case=False,
            max_results=120,
        )

        if result == "没有找到匹配内容。":
            return "未发现 TODO/FIXME/BUG/HACK/XXX 标记。"

        return "待办与风险标记报告：\n" + result

    except Exception as e:
        return f"生成待办报告失败：{e}"


def python_symbols(root_dir: str, path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        files = [target] if target.is_file() else _iter_files(target)
        blocks = []

        for file_path in files:
            if file_path.suffix.lower() != ".py":
                continue

            try:
                source = file_path.read_text(encoding="utf-8", errors="ignore")
                tree = ast.parse(source)
            except Exception as e:
                blocks.append(f"[{file_path.relative_to(root)}]\n解析失败：{e}")
                continue

            symbols = []
            imports = []

            for node in ast.iter_child_nodes(tree):
                if isinstance(node, ast.ClassDef):
                    symbols.append(f"- class {node.name} (line {node.lineno})")
                elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    prefix = "async def" if isinstance(node, ast.AsyncFunctionDef) else "def"
                    symbols.append(f"- {prefix} {node.name} (line {node.lineno})")
                elif isinstance(node, ast.Import):
                    imports.extend(alias.name for alias in node.names)
                elif isinstance(node, ast.ImportFrom):
                    module = node.module or ""
                    imports.append(f"from {module}")

            if symbols or imports:
                import_text = ", ".join(sorted(set(imports))[:20]) or "无顶层 import"
                symbol_text = "\n".join(symbols) or "无顶层类/函数"
                blocks.append(
                    f"[{file_path.relative_to(root)}]\n"
                    f"Imports: {import_text}\n"
                    f"Symbols:\n{symbol_text}"
                )

            if len(blocks) >= 40:
                blocks.append("[结果过多，已截断]")
                break

        return "\n\n".join(blocks) if blocks else "未发现 Python 符号。"

    except Exception as e:
        return f"提取 Python 符号失败：{e}"


def dependency_report(root_dir: str, path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        files = [target] if target.is_file() else _iter_files(target)
        reports = []

        for file_path in files:
            name = file_path.name.lower()
            relative = file_path.relative_to(root)

            if name == "requirements.txt":
                deps = []
                for line in file_path.read_text(encoding="utf-8", errors="ignore").splitlines():
                    stripped = line.strip()
                    if stripped and not stripped.startswith("#"):
                        deps.append(stripped)
                reports.append(f"[{relative}]\n" + ("\n".join(f"- {dep}" for dep in deps) or "无依赖"))

            elif name == "pyproject.toml":
                data = tomllib.loads(file_path.read_text(encoding="utf-8", errors="ignore"))
                deps = data.get("project", {}).get("dependencies", [])
                optional = data.get("project", {}).get("optional-dependencies", {})
                poetry_deps = data.get("tool", {}).get("poetry", {}).get("dependencies", {})
                lines = [f"- {dep}" for dep in deps]
                lines.extend(f"- {name}: {values}" for name, values in optional.items())
                lines.extend(f"- {name}: {version}" for name, version in poetry_deps.items() if name.lower() != "python")
                reports.append(f"[{relative}]\n" + ("\n".join(lines) or "未发现依赖字段"))

            elif name == "package.json":
                data = json.loads(file_path.read_text(encoding="utf-8", errors="ignore"))
                deps = data.get("dependencies", {})
                dev_deps = data.get("devDependencies", {})
                scripts = data.get("scripts", {})
                lines = ["dependencies:"]
                lines.extend(f"- {name}: {version}" for name, version in deps.items())
                lines.append("devDependencies:")
                lines.extend(f"- {name}: {version}" for name, version in dev_deps.items())
                lines.append("scripts:")
                lines.extend(f"- {name}: {command}" for name, command in scripts.items())
                reports.append(f"[{relative}]\n" + "\n".join(lines))

        return "\n\n".join(reports) if reports else "未发现 requirements.txt、pyproject.toml 或 package.json。"

    except Exception as e:
        return f"生成依赖报告失败：{e}"


def compare_files(root_dir: str, left_path: str, right_path: str, context_lines: int = 3) -> str:
    try:
        left = _safe_path(root_dir, left_path)
        right = _safe_path(root_dir, right_path)

        for target, label in [(left, left_path), (right, right_path)]:
            if not target.exists():
                return f"文件不存在：{label}"
            if not target.is_file():
                return f"不是文件：{label}"
            if not _is_text_file(target):
                return f"当前只支持比较纯文本文件，不支持：{label}"

        left_lines = left.read_text(encoding="utf-8", errors="ignore").splitlines(keepends=True)
        right_lines = right.read_text(encoding="utf-8", errors="ignore").splitlines(keepends=True)
        diff = difflib.unified_diff(
            left_lines,
            right_lines,
            fromfile=left_path,
            tofile=right_path,
            n=max(0, min(context_lines, 10)),
        )
        text = "".join(diff)
        return text[:MAX_TOTAL_CHARS] + ("\n[diff 过长，已截断]" if len(text) > MAX_TOTAL_CHARS else "") or "两个文件内容一致。"

    except Exception as e:
        return f"比较文件失败：{e}"


def preview_replace_in_file(root_dir: str, path: str, old_text: str, new_text: str) -> str:
    try:
        target = _safe_path(root_dir, path)

        if not target.exists():
            return f"文件不存在：{path}"

        if not target.is_file():
            return f"不是文件：{path}"

        if not _is_text_file(target):
            return f"当前只支持预览纯文本文件替换，不支持：{target.suffix}"

        original = target.read_text(encoding="utf-8", errors="ignore")

        if old_text not in original:
            return f"未找到要替换的内容：{old_text}"

        updated = original.replace(old_text, new_text)
        diff = difflib.unified_diff(
            original.splitlines(keepends=True),
            updated.splitlines(keepends=True),
            fromfile=path,
            tofile=f"{path} (preview)",
            n=3,
        )
        text = "".join(diff)
        return text[:MAX_TOTAL_CHARS] + ("\n[diff 过长，已截断]" if len(text) > MAX_TOTAL_CHARS else "")

    except Exception as e:
        return f"预览替换失败：{e}"


def search_file(root_dir: str, keyword: str, path: str = ".") -> str:
    try:
        target = _safe_path(root_dir, path)
        root = _resolve_root(root_dir)

        if not target.exists():
            return f"路径不存在：{path}"

        results = []
        files = [target] if target.is_file() else _iter_supported_files(target)

        for file_path in files:
            suffix = file_path.suffix.lower()

            if suffix in {".pdf", ".docx"}:
                content = read_file(
                    root_dir=root_dir,
                    path=str(file_path.relative_to(root)),
                    max_chars=MAX_FILE_CHARS,
                )
                lines = content.splitlines()
            else:
                try:
                    lines = file_path.read_text(encoding="utf-8", errors="ignore").splitlines()
                except Exception:
                    continue

            for line_no, line in enumerate(lines, start=1):
                if keyword.lower() in line.lower():
                    relative = file_path.relative_to(root)
                    results.append(f"{relative}:{line_no}: {line.strip()}")

                    if len(results) >= MAX_SEARCH_RESULTS:
                        return "\n".join(results) + "\n\n[搜索结果过多，已截断]"

        return "\n".join(results) if results else "没有找到匹配内容。"

    except Exception as e:
        return f"搜索失败：{e}"


def write_file(root_dir: str, path: str, content: str, overwrite: bool = False) -> str:
    """
    写入新文件或覆盖已有文件。支持纯文本文件和 .docx。
    默认不覆盖，避免意外破坏已有文件。
    """
    try:
        target = _safe_path(root_dir, path)

        if target.exists() and not overwrite:
            return f"文件已存在：{path}。如需覆盖，请设置 overwrite=true。"

        if not _is_supported_write_file(target):
            return f"暂不支持写入该类型文件：{target.suffix or '[无扩展名]'}。为避免生成损坏文件，已取消写入。"

        target.parent.mkdir(parents=True, exist_ok=True)

        backup_msg = ""
        if target.exists():
            backup_path = _backup_file(target)
            backup_msg = f"\n已备份原文件：{backup_path.name}"

        if _is_docx_file(target):
            _write_docx(target, content)
        else:
            target.write_text(content, encoding="utf-8")

        return f"写入成功：{path}{backup_msg}"

    except Exception as e:
        return f"写入文件失败：{e}"


def replace_in_file(root_dir: str, path: str, old_text: str, new_text: str) -> str:
    """
    在纯文本文件或 .docx 文件中替换内容。
    修改前自动备份。
    """
    try:
        target = _safe_path(root_dir, path)

        if not target.exists():
            return f"文件不存在：{path}"

        if not target.is_file():
            return f"不是文件：{path}"

        if not _is_supported_write_file(target):
            return f"当前只支持修改纯文本文件和 .docx，不支持直接修改：{target.suffix}"

        if _is_docx_file(target):
            content = _read_docx(target, MAX_FILE_CHARS)

            if old_text not in content:
                return f"未找到要替换的内容：{old_text}"

            backup_path = _backup_file(target)
            changed = _replace_docx_text(target, old_text, new_text)

            if not changed:
                return f"未找到要替换的内容：{old_text}"

            return (
                f"替换成功：{path}\n"
                f"已备份原文件：{backup_path.name}\n"
                f"替换内容：{old_text} -> {new_text}"
            )

        text = target.read_text(encoding="utf-8", errors="ignore")

        if old_text not in text:
            return f"未找到要替换的内容：{old_text}"

        backup_path = _backup_file(target)
        new_content = text.replace(old_text, new_text)

        target.write_text(new_content, encoding="utf-8")

        return (
            f"替换成功：{path}\n"
            f"已备份原文件：{backup_path.name}\n"
            f"替换内容：{old_text} -> {new_text}"
        )

    except Exception as e:
        return f"替换文件内容失败：{e}"


def append_file(root_dir: str, path: str, content: str) -> str:
    """
    向纯文本文件或 .docx 文件追加内容。
    修改前自动备份。
    """
    try:
        target = _safe_path(root_dir, path)

        if not _is_supported_write_file(target):
            return f"当前只支持追加纯文本文件和 .docx，不支持：{target.suffix or '[无扩展名]'}"

        target.parent.mkdir(parents=True, exist_ok=True)

        backup_msg = ""
        if target.exists():
            backup_path = _backup_file(target)
            backup_msg = f"\n已备份原文件：{backup_path.name}"

        if _is_docx_file(target):
            _append_docx(target, content)
        else:
            with target.open("a", encoding="utf-8") as f:
                f.write(content)

        return f"追加成功：{path}{backup_msg}"

    except Exception as e:
        return f"追加文件失败：{e}"
