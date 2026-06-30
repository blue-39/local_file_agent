import hashlib
import os
import re
from pathlib import Path

from dotenv import load_dotenv
from langchain_core.documents import Document


load_dotenv()

os.environ.setdefault("HF_HUB_DISABLE_SYMLINKS_WARNING", "1")

DEFAULT_EMBEDDING_MODEL = os.getenv("RAG_EMBEDDING_MODEL", "BAAI/bge-small-zh-v1.5")
DEFAULT_PERSIST_DIRECTORY = "chroma_db"
DEFAULT_MODEL_CACHE_DIR = os.getenv("RAG_MODEL_CACHE_DIR") or None
ADD_DOCUMENTS_BATCH_SIZE = 32

SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".toml",
    ".csv", ".log", ".ini", ".cfg", ".html", ".css", ".js",
    ".ts", ".tsx", ".jsx", ".java", ".cpp", ".c", ".h", ".hpp",
    ".go", ".rs", ".sh", ".sql", ".pdf", ".docx",
}

SKIP_DIRS = {
    ".git", ".hg", ".svn", "__pycache__", ".pytest_cache", ".mypy_cache",
    ".ruff_cache", ".idea", ".vscode", "node_modules", ".venv", "venv",
    "dist", "build", ".next", ".turbo", "chroma_db",
}


def _resolve_root(root_dir: str) -> Path:
    root = Path(root_dir).expanduser().resolve()

    if not root.exists():
        raise ValueError(f"根目录不存在：{root_dir}")

    if not root.is_dir():
        raise ValueError(f"不是目录：{root_dir}")

    return root


def _safe_path(root_dir: str, path: str = ".") -> Path:
    root = _resolve_root(root_dir)
    target = (root / path).expanduser().resolve()

    try:
        target.relative_to(root)
    except ValueError:
        raise ValueError("非法路径：只能访问用户指定的根目录内部。")

    return target


def _iter_files(target: Path):
    if target.is_file():
        yield target
        return

    for file_path in target.rglob("*"):
        if any(part in SKIP_DIRS for part in file_path.parts):
            continue

        if file_path.is_file():
            yield file_path


def _read_pdf(path: Path, max_chars: int) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        chunks = []

        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                chunks.append(text.strip())

            if sum(len(item) for item in chunks) >= max_chars:
                break

        return "\n\n".join(chunks)[:max_chars]

    except Exception as e:
        raise RuntimeError(f"读取 PDF 失败：{e}") from e


def _read_docx(path: Path, max_chars: int) -> str:
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        chunks = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    chunks.append(row_text)

        return "\n".join(chunks)[:max_chars]

    except Exception as e:
        raise RuntimeError(f"读取 DOCX 失败：{e}") from e


def _read_document_file(path: Path, max_chars: int) -> str:
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _read_pdf(path, max_chars=max_chars)

    if suffix == ".docx":
        return _read_docx(path, max_chars=max_chars)

    return path.read_text(encoding="utf-8", errors="ignore")[:max_chars]


def _clean_text(text) -> str:
    if text is None:
        return ""

    if not isinstance(text, str):
        text = str(text)

    text = text.replace("\x00", " ")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _sanitize_documents(documents: list[Document]) -> list[Document]:
    sanitized = []

    for document in documents:
        text = _clean_text(document.page_content)

        if not text:
            continue

        metadata = {
            "source": str(document.metadata.get("source", "")),
            "root_dir": str(document.metadata.get("root_dir", "")),
            "suffix": str(document.metadata.get("suffix", "")),
        }
        sanitized.append(Document(page_content=text, metadata=metadata))

    return sanitized


def _add_documents_in_batches(vectorstore, chunks: list[Document]) -> None:
    for start in range(0, len(chunks), ADD_DOCUMENTS_BATCH_SIZE):
        batch = chunks[start:start + ADD_DOCUMENTS_BATCH_SIZE]

        try:
            vectorstore.add_documents(batch)
        except Exception as e:
            sources = sorted({doc.metadata.get("source", "[unknown]") for doc in batch})
            raise RuntimeError(
                "写入向量库失败。可能的文件来源："
                f"{', '.join(sources[:8])}。原始错误：{e}"
            ) from e


def get_collection_name(root_dir: str) -> str:
    """Return a stable Chroma collection name for a root directory."""
    root = str(_resolve_root(root_dir))
    safe = re.sub(r"[^A-Za-z0-9_-]+", "_", root).strip("_").lower()
    digest = hashlib.sha1(root.encode("utf-8")).hexdigest()[:12]
    base = f"agent_{safe}_{digest}" if safe else f"agent_{digest}"

    if len(base) > 63:
        base = f"{base[:50]}_{digest}"

    return base


def get_embeddings():
    """Create the local HuggingFace embedding model."""
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except ImportError as e:
        raise ImportError(
            "缺少 RAG 依赖 langchain-huggingface / sentence-transformers，"
            "请先执行：pip install -r requirements.txt"
        ) from e

    return HuggingFaceEmbeddings(
        model_name=DEFAULT_EMBEDDING_MODEL,
        cache_folder=DEFAULT_MODEL_CACHE_DIR,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def get_vectorstore(root_dir: str):
    """Open the Chroma vector store for a root directory."""
    try:
        from langchain_chroma import Chroma
    except ImportError as e:
        raise ImportError(
            "缺少 RAG 依赖 langchain-chroma，请先执行：pip install -r requirements.txt"
        ) from e

    root = _resolve_root(root_dir)
    persist_directory = str(root / DEFAULT_PERSIST_DIRECTORY)

    return Chroma(
        collection_name=get_collection_name(root_dir),
        embedding_function=get_embeddings(),
        persist_directory=persist_directory,
    )


def load_documents_from_dir(
    root_dir: str,
    path: str = ".",
    max_files: int = 200,
    max_chars_each: int = 20000,
) -> list[Document]:
    """Load supported documents from root_dir/path into LangChain Document objects."""
    root = _resolve_root(root_dir)
    target = _safe_path(root_dir, path)

    if not target.exists():
        raise ValueError(f"路径不存在：{path}")

    documents = []

    for file_path in _iter_files(target):
        if file_path.suffix.lower() not in SUPPORTED_DOCUMENT_EXTENSIONS:
            continue

        if len(documents) >= max_files:
            break

        relative = str(file_path.relative_to(root))

        try:
            text = _clean_text(_read_document_file(file_path, max_chars=max_chars_each))
        except Exception:
            continue

        if not text:
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": relative,
                    "root_dir": str(root),
                    "suffix": file_path.suffix.lower(),
                },
            )
        )

    return documents


def build_or_update_index(
    root_dir: str,
    path: str = ".",
    chunk_size: int = 800,
    chunk_overlap: int = 120,
    max_files: int = 200,
) -> str:
    """Build or update the local Chroma index for files under root_dir/path."""
    try:
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
        except ImportError as e:
            raise ImportError(
                "缺少 RAG 依赖 langchain-text-splitters，请先执行：pip install -r requirements.txt"
            ) from e

        documents = load_documents_from_dir(
            root_dir=root_dir,
            path=path,
            max_files=max_files,
        )

        if not documents:
            return "索引构建取消：没有加载到可索引文档。"

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", "。", "；", "，", " ", ""],
        )
        chunks = _sanitize_documents(splitter.split_documents(documents))

        if not chunks:
            return "索引构建取消：文档切分后没有生成 chunk。"

        vectorstore = get_vectorstore(root_dir)
        _add_documents_in_batches(vectorstore, chunks)

        persist = getattr(vectorstore, "persist", None)
        if callable(persist):
            persist()

        root = _resolve_root(root_dir)
        return (
            "索引构建完成。\n"
            f"文档数：{len(documents)}\n"
            f"chunk 数：{len(chunks)}\n"
            f"collection 名称：{get_collection_name(root_dir)}\n"
            f"向量库目录：{root / DEFAULT_PERSIST_DIRECTORY}"
        )

    except Exception as e:
        return f"建立 RAG 索引失败：{e}"


def semantic_search(root_dir: str, query: str, k: int = 5) -> str:
    """Run semantic search against the local Chroma index for root_dir."""
    try:
        query = query.strip()

        if not query:
            return "语义检索失败：query 不能为空。"

        k = max(1, min(k, 20))
        vectorstore = get_vectorstore(root_dir)
        results = vectorstore.similarity_search_with_score(query, k=k)

        if not results:
            return "没有检索到结果。请先建立索引，或扩大索引目录后重试。"

        blocks = []
        for index, (document, score) in enumerate(results, start=1):
            content = document.page_content.strip()
            if len(content) > 2000:
                content = content[:2000] + "\n[content 过长，已截断]"

            source = document.metadata.get("source", "[unknown]")
            blocks.append(
                f"Result {index}\n"
                f"source: {source}\n"
                f"score: {score}\n"
                f"content:\n{content}"
            )

        return "\n\n" + ("-" * 60 + "\n\n").join(blocks)

    except Exception as e:
        return f"语义检索失败：{e}。如果尚未建立索引，请先调用 index_directory。"
